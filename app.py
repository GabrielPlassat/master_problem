"""
Formulateur de Sous-Problème Stratégique – Transition Écologique
Outil Streamlit + Claude API
"""

import streamlit as st
import anthropic
import json
import io
from datetime import datetime

# ─── CONFIG PAGE ────────────────────────────────────────────────────────────────
st.set_page_config(
    page_title="Formulateur Stratégique – Transition Écologique",
    page_icon="🌿",
    layout="wide",
    initial_sidebar_state="expanded"
)

# ─── CLÉ API — chargée depuis st.secrets ────────────────────────────────────────
def get_api_key() -> str:
    """Charge la clé API depuis st.secrets (secrets.toml ou Streamlit Cloud)."""
    try:
        return st.secrets["ANTHROPIC_API_KEY"]
    except (KeyError, FileNotFoundError):
        return ""

API_KEY = get_api_key()

# ─── STYLES ─────────────────────────────────────────────────────────────────────
st.markdown("""
<style>
.phase-header {
    background: linear-gradient(135deg, #1a7a4a, #2d9d6a);
    padding: 20px 28px;
    border-radius: 12px;
    color: white;
    margin-bottom: 24px;
}
.phase-header h2 { color: white; margin: 0 0 6px 0; }
.phase-header p  { color: #d4f5e9; margin: 0; }
.hmw-box {
    background: linear-gradient(135deg, #0f4c75, #1b6ca8);
    padding: 28px;
    border-radius: 14px;
    color: white;
    text-align: center;
    margin: 20px 0;
}
.hmw-box h3 { color: white; font-size: 1.4rem; margin: 0; line-height: 1.5; }
.score-box {
    border: 2px solid #e0e0e0;
    border-radius: 10px;
    padding: 16px;
    text-align: center;
}
</style>
""", unsafe_allow_html=True)

# ─── CONSTANTES ─────────────────────────────────────────────────────────────────
PHASES = [
    ("1", "Ancrage",                  "Comprendre et clarifier le problème général"),
    ("2", "Miroir des Impossibles",   "Déstabiliser les certitudes de l'expert"),
    ("3", "Scout Technologique",      "Explorer les solutions venues d'ailleurs"),
    ("4", "Décomposition",            "Identifier les sous-problèmes candidats"),
    ("5", "Ressources & Contraintes", "Décrire précisément vos moyens"),
    ("6", "Test de Cohérence",        "Évaluer la faisabilité du sous-problème"),
    ("7", "Formulation Finale",       "Formuler, exporter et agir"),
]

CYNEFIN_INFO = {
    "Clair":      ("🟢", "Relations cause-effet connues. Appliquer les meilleures pratiques."),
    "Compliqué":  ("🟡", "Plusieurs bonnes réponses. Requiert expertise et analyse."),
    "Complexe":   ("🟠", "Imprévisible. Expérimenter, observer, adapter."),
    "Pernicieux": ("🔴", "Pas de solution définitive. Problème social profond (Wicked Problem)."),
}

# ─── SESSION STATE ───────────────────────────────────────────────────────────────
DEFAULTS = {
    "phase": 0,
    "doc_text": "",
    "sofia_data": None,
    "problem_input": "",
    "reformulation": "",
    "cynefin": "",
    "ambiguites": [],
    "manques": [],
    # Phase 2 — Miroir des Impossibles
    "miroir_cas": [],           # liste des cas générés
    "miroir_reactions": {},     # réactions de l'expert par cas
    "miroir_lecon": "",         # leçon tirée par l'expert
    # Phase 3 — Scout Technologique
    "scout_techs": [],          # technologies candidates générées
    "scout_selections": [],     # technologies retenues par l'expert
    "scout_ouvertures": "",     # reformulation après le scout
    # Phase 4 — Décomposition
    "sub_problems": [],
    "selected_sub_problems_list": [],
    "selected_sub_problem": "",
    # Phases suivantes
    "resources": {},
    "constraints": {},
    "coherence_score": None,
    "coherence_details": {},
    "hmw": "",
    "final_data": {},
}
for k, v in DEFAULTS.items():
    if k not in st.session_state:
        st.session_state[k] = v

# ─── HELPERS ────────────────────────────────────────────────────────────────────

def parse_sofia_html(html_content: str) -> dict:
    """
    Parse un rapport HTML exporté par SofIA (sofia-transition-ecologique.fr).
    Retourne un dict structuré avec questions, reformulations, réponses et sources ADEME.
    """
    from html.parser import HTMLParser
    import re

    # Utilise BeautifulSoup si disponible, sinon regex fallback
    try:
        from bs4 import BeautifulSoup
        soup = BeautifulSoup(html_content, "html.parser")

        # Date de la conversation
        h1 = soup.find("h1")
        date_conv = h1.get_text(strip=True) if h1 else ""

        exchanges = []

        # Questions
        questions = soup.find_all("div", class_="question")
        reformulations = soup.find_all("div", class_="reformulation")
        responses = soup.find_all("div", class_="response")
        sources_blocks = soup.find_all("div", class_="sources")

        for i, (q, resp) in enumerate(zip(questions, responses)):
            # Question brute
            q_text = q.get_text(separator=" ", strip=True)
            q_text = re.sub(r"^\d+\.\s*Question\s*:", "", q_text).strip()

            # Reformulation SofIA
            ref_text = ""
            if i < len(reformulations):
                ref_text = reformulations[i].get_text(separator=" ", strip=True)
                ref_text = re.sub(r"Reformulation de la question\s*:", "", ref_text).strip()

            # Réponse SofIA (texte enrichi)
            # On garde le texte principal, on supprime les balises h2 "Sofia:"
            for tag in resp.find_all("h2"):
                tag.decompose()
            resp_text = resp.get_text(separator="\n", strip=True)
            # Nettoyage des artefacts HTML
            resp_text = re.sub(r"\n{3,}", "\n\n", resp_text)

            # Sources ADEME associées
            sources = []
            if i < len(sources_blocks):
                cards = sources_blocks[i].find_all("div", class_="source-card")
                for card in cards:
                    card_id = card.get("id", "")
                    title_tag = card.find("h2", class_="card-title")
                    title = title_tag.get_text(strip=True) if title_tag else ""
                    # Lien PDF
                    link_tag = card.find("a", href=True)
                    link = link_tag["href"] if link_tag else ""
                    # Texte extrait
                    text_tag = card.find("p", class_="card-text")
                    excerpt = text_tag.get_text(separator=" ", strip=True) if text_tag else ""
                    # Score de pertinence
                    score_tag = card.find("p", class_="card-similarity")
                    score_str = score_tag.get_text(strip=True) if score_tag else ""
                    score_val = re.search(r"([\d.]+)\s*%", score_str)
                    score_pct = float(score_val.group(1)) if score_val else 0.0

                    sources.append({
                        "id": card_id,
                        "titre": title,
                        "lien": link,
                        "extrait": excerpt[:400],
                        "score_pertinence": score_pct,
                    })
                # Tri par score décroissant
                sources.sort(key=lambda x: x["score_pertinence"], reverse=True)

            exchanges.append({
                "question": q_text,
                "reformulation_sofia": ref_text,
                "reponse_sofia": resp_text[:3000],  # Cap pour éviter token overflow
                "nb_sources": len(sources),
                "sources_top5": sources[:5],
            })

        return {
            "is_sofia": True,
            "date": date_conv,
            "nb_echanges": len(exchanges),
            "exchanges": exchanges,
        }

    except ImportError:
        # Fallback regex si BeautifulSoup absent
        import re
        questions = re.findall(r'class="question"[^>]*>(.*?)</div>', html_content, re.DOTALL)
        responses = re.findall(r'class="response"[^>]*>(.*?)</div>', html_content, re.DOTALL)
        clean = lambda s: re.sub(r"<[^>]+>", " ", s).strip()
        exchanges = [
            {"question": clean(q), "reponse_sofia": clean(r)[:2000], "sources_top5": []}
            for q, r in zip(questions, responses)
        ]
        return {"is_sofia": True, "date": "", "nb_echanges": len(exchanges), "exchanges": exchanges}


def sofia_to_context(sofia_data: dict) -> str:
    """Convertit les données SofIA en bloc de contexte structuré pour Claude."""
    lines = [
        f"=== RAPPORT SOFIA — Transition Écologique ===",
        f"Date : {sofia_data.get('date', '')}",
        f"Nombre d'échanges : {sofia_data.get('nb_echanges', 0)}",
        "",
    ]
    for i, ex in enumerate(sofia_data.get("exchanges", []), 1):
        lines.append(f"── Échange {i} ──")
        lines.append(f"QUESTION : {ex.get('question', '')}")
        if ex.get("reformulation_sofia"):
            lines.append(f"REFORMULATION SOFIA : {ex.get('reformulation_sofia', '')}")
        lines.append(f"RÉPONSE SOFIA (synthèse ADEME) :\n{ex.get('reponse_sofia', '')}")
        top_sources = ex.get("sources_top5", [])
        if top_sources:
            lines.append(f"\nSOURCES ADEME LES PLUS PERTINENTES ({len(top_sources)}) :")
            for s in top_sources:
                lines.append(
                    f"  • [{s.get('score_pertinence', 0):.1f}%] {s.get('titre', '')} — {s.get('extrait', '')[:200]}"
                )
        lines.append("")
    return "\n".join(lines)


def extract_text(uploaded_file) -> str:
    """Extrait le texte d'un fichier PDF, DOCX, TXT ou HTML SofIA."""
    ext = uploaded_file.name.rsplit(".", 1)[-1].lower()
    raw_bytes = uploaded_file.read()
    try:
        if ext == "txt":
            return raw_bytes.decode("utf-8", errors="ignore")
        elif ext == "pdf":
            import pypdf
            reader = pypdf.PdfReader(io.BytesIO(raw_bytes))
            return "\n".join(p.extract_text() or "" for p in reader.pages)
        elif ext in ("docx", "doc"):
            from docx import Document
            doc = Document(io.BytesIO(raw_bytes))
            return "\n".join(p.text for p in doc.paragraphs)
        elif ext in ("html", "htm"):
            html_str = raw_bytes.decode("utf-8", errors="ignore")
            # Détecte si c'est un rapport SofIA
            if "sofia" in html_str.lower()[:2000] or "class=\"question\"" in html_str or "card-similarity" in html_str:
                sofia_data = parse_sofia_html(html_str)
                # Stocke aussi les données structurées pour affichage
                st.session_state["sofia_data"] = sofia_data
                return sofia_to_context(sofia_data)
            else:
                # HTML générique : strip balises
                import re
                return re.sub(r"<[^>]+>", " ", html_str)
    except Exception as e:
        st.error(f"Erreur lecture {uploaded_file.name} : {e}")
    return ""


def claude(prompt: str, system: str = None, max_tokens: int = 2000) -> str | None:
    """Appel Claude API avec gestion d'erreur."""
    if not API_KEY:
        st.error("⚠️ Clé API manquante. Ajoutez ANTHROPIC_API_KEY dans .streamlit/secrets.toml")
        return None
    sys = system or (
        "Tu es un consultant senior spécialisé en transition écologique et formulation "
        "de problèmes stratégiques. Tu réponds en français, avec précision et pédagogie."
    )
    try:
        client = anthropic.Anthropic(api_key=API_KEY)
        msg = client.messages.create(
            model="claude-sonnet-4-6",
            max_tokens=max_tokens,
            system=sys,
            messages=[{"role": "user", "content": prompt}],
        )
        return msg.content[0].text
    except Exception as e:
        st.error(f"Erreur API : {e}")
        return None


def parse_json(raw: str) -> dict | list | None:
    """Parse JSON en nettoyant les éventuels blocs markdown."""
    if not raw:
        return None
    clean = raw.strip().removeprefix("```json").removeprefix("```").removesuffix("```").strip()
    try:
        return json.loads(clean)
    except Exception as e:
        st.error(f"Erreur de parsing JSON : {e}\n\nRéponse brute : {raw[:300]}")
        return None


def score_emoji(s: int, max_s: int = 100) -> str:
    pct = s / max_s * 100
    return "🟢" if pct >= 70 else "🟡" if pct >= 40 else "🔴"


def next_phase():
    st.session_state.phase += 1
    st.rerun()


def go_phase(n: int):
    st.session_state.phase = n
    st.rerun()


# ─── SIDEBAR ────────────────────────────────────────────────────────────────────
with st.sidebar:
    st.markdown("## 🌿 Formulateur Stratégique")
    st.caption("Transition Écologique — Outil de cadrage")

    # Indicateur de statut API (discret)
    if API_KEY:
        st.caption("🟢 API connectée")
    else:
        st.error("🔴 Clé API manquante\n\nAjoutez `ANTHROPIC_API_KEY` dans `.streamlit/secrets.toml`")

    st.divider()
    st.markdown("### 📍 Progression")
    for i, (num, name, _) in enumerate(PHASES):
        if i < st.session_state.phase:
            st.markdown(f"✅ **Phase {num}** – {name}")
        elif i == st.session_state.phase:
            st.markdown(f"▶️ **Phase {num} – {name}** ← ici")
        else:
            st.markdown(f"⬜ Phase {num} – {name}")

    st.divider()
    col1, col2 = st.columns(2)
    with col1:
        if st.button("◀ Retour", disabled=st.session_state.phase == 0, use_container_width=True):
            st.session_state.phase -= 1
            st.rerun()
    with col2:
        if st.button("🔄 Reset", use_container_width=True):
            for k in list(st.session_state.keys()):
                del st.session_state[k]
            for k, v in DEFAULTS.items():
                st.session_state[k] = v
            st.rerun()

    st.divider()
    st.caption("Basé sur : Cynefin (Snowden), Problem Statement Canvas, JTBD, Double Diamant, HMW (IDEO), Frame Innovation (Dorst)")


# ─── PHASE HEADER HELPER ────────────────────────────────────────────────────────
def phase_header(num, name, desc):
    st.markdown(f"""
<div class="phase-header">
  <h2>Phase {num} — {name}</h2>
  <p>{desc}</p>
</div>
""", unsafe_allow_html=True)


# ════════════════════════════════════════════════════════════════════════════════
# PHASE 0 — ANCRAGE
# ════════════════════════════════════════════════════════════════════════════════
if st.session_state.phase == 0:
    phase_header("1", "Ancrage", "Décrire et clarifier votre problème général. Briser l'illusion de clarté.")

    st.markdown("""
> *"Tomber amoureux du problème, pas de la solution."* — Marty Cagan
>
> Beaucoup d'experts pensent avoir un problème clair. Ce qu'ils ont, c'est souvent **une intention**.
> Cette phase va révéler les zones d'ombre avant d'aller plus loin.
""")

    # ── Upload documents ──
    st.subheader("📎 Documents de contexte")

    col_tip1, col_tip2 = st.columns(2)
    with col_tip1:
        st.info(
            "**💡 Rapport SofIA recommandé**\n\n"
            "Consultez d'abord [SofIA](https://www.sofia-transition-ecologique.fr) "
            "(IA de l'ADEME) pour interroger la base documentaire sur votre thématique, "
            "puis exportez l'historique en HTML et uploadez-le ici.\n\n"
            "L'outil reconnaît automatiquement ce format et extrait questions, "
            "réponses et sources ADEME."
        )
    with col_tip2:
        st.markdown(
            "**Autres formats acceptés :**\n"
            "- 📄 PDF (rapports, études)\n"
            "- 📝 DOCX (cahiers des charges)\n"
            "- 🌐 HTML (rapport SofIA exporté)\n"
            "- 📋 TXT (notes, synthèses)\n"
        )

    files = st.file_uploader(
        "Glissez vos documents ici",
        accept_multiple_files=True,
        type=["pdf", "docx", "txt", "html", "htm"],
        label_visibility="collapsed"
    )
    if files:
        texts = []
        sofia_detected = False
        for f in files:
            t = extract_text(f)
            if t:
                texts.append(f"=== {f.name} ===\n{t}")
            if st.session_state.get("sofia_data"):
                sofia_detected = True

        st.session_state.doc_text = "\n\n".join(texts)
        nb_chars = len(st.session_state.doc_text)

        if sofia_detected and st.session_state.sofia_data:
            sd = st.session_state.sofia_data
            st.success(f"✅ Rapport SofIA détecté — {sd['nb_echanges']} échange(s) analysé(s)")
            st.markdown(f"**{sd.get('date', '')}**")

            for i, ex in enumerate(sd.get("exchanges", []), 1):
                with st.expander(f"💬 Échange {i} : {ex.get('question', '')[:80]}…"):
                    st.markdown(f"**Question :** {ex.get('question', '')}")
                    if ex.get("reformulation_sofia"):
                        st.markdown(f"**Reformulation SofIA :** *{ex.get('reformulation_sofia', '')}*")
                    st.markdown("**Réponse SofIA (extrait) :**")
                    st.markdown(ex.get("reponse_sofia", "")[:1500] + "…")

                    top_src = ex.get("sources_top5", [])
                    if top_src:
                        st.markdown(f"**📚 {ex.get('nb_sources', len(top_src))} sources ADEME — Top 5 :**")
                        for s in top_src:
                            score_color = "🟢" if s["score_pertinence"] > 10 else "🟡" if s["score_pertinence"] > 5 else "🔴"
                            lien_md = f"[↗]({s['lien']})" if s.get("lien") else ""
                            st.markdown(
                                f"  {score_color} **{s['score_pertinence']:.1f}%** — {s.get('titre', '')} {lien_md}"
                            )
        else:
            st.success(f"✅ {len(files)} document(s) chargé(s) — {nb_chars:,} caractères extraits")
            with st.expander("Aperçu du texte extrait"):
                st.text(st.session_state.doc_text[:3000] + ("…" if nb_chars > 3000 else ""))

    st.divider()

    # ── Saisie libre ──
    st.subheader("✍️ Décrivez votre problème général")
    st.caption("Exprimez-vous librement. Ne pensez pas encore à la solution ni aux ressources.")
    problem_input = st.text_area(
        "Problème général",
        value=st.session_state.problem_input,
        height=160,
        placeholder=(
            "Ex : Les émissions de GES liées à la logistique urbaine continuent d'augmenter "
            "malgré les réglementations. Les opérateurs peinent à adapter leurs flottes. "
            "Les collectivités manquent de données pour agir efficacement..."
        ),
        label_visibility="collapsed"
    )
    st.session_state.problem_input = problem_input

    can_analyze = bool(problem_input.strip()) and bool(API_KEY)

    if not API_KEY:
        st.warning("⚠️ Renseignez votre clé API Anthropic dans la barre latérale.")

    if st.button("🔍 Analyser mon problème", type="primary", disabled=not can_analyze):
        # Contexte : priorité au rapport SofIA (données ADEME structurées)
        context_block = ""
        if st.session_state.get("sofia_data"):
            sd = st.session_state.sofia_data
            sofia_lines = [
                "CONTEXTE ISSU DE SOFIA (base documentaire ADEME — source de référence) :",
                f"Date de consultation : {sd.get('date', '')}",
                "",
            ]
            for i, ex in enumerate(sd.get("exchanges", []), 1):
                sofia_lines.append(f"Q{i} posée par l'expert : {ex.get('question', '')}")
                sofia_lines.append(f"Synthèse ADEME : {ex.get('reponse_sofia', '')[:800]}")
                top5 = ex.get("sources_top5", [])
                if top5:
                    sofia_lines.append("Sources ADEME mobilisées : " +
                        " | ".join(f"{s['titre'][:50]} ({s['score_pertinence']:.0f}%)" for s in top5[:3]))
                sofia_lines.append("")
            context_block = "\n".join(sofia_lines)[:5000] + "\n\n"
        elif st.session_state.doc_text:
            context_block = f"DOCUMENTS FOURNIS :\n{st.session_state.doc_text[:4000]}\n\n"

        sofia_instruction = ""
        if st.session_state.get("sofia_data"):
            sofia_instruction = (
                "\nATTENTION : Le contexte provient de SofIA (IA de l'ADEME). "
                "Utilise les données chiffrées, objectifs réglementaires et retours d'expérience "
                "identifiés dans les échanges SofIA pour enrichir ton analyse et tes recommandations. "
                "Cite les données concrètes issues des sources ADEME pour justifier ta classification Cynefin."
            )

        prompt = f"""{context_block}L'expert décrit son problème général :
"{problem_input}"{sofia_instruction}

Réalise une analyse structurée. Réponds UNIQUEMENT en JSON valide (pas de markdown, pas de backticks) avec ces clés :
- "reformulation" : reformulation précise en 2-3 phrases, fidèle à l'expert mais sans ambiguïté. Si des données SofIA/ADEME sont disponibles, intègre les éléments factuels clés (objectifs, chiffres, contexte réglementaire).
- "cynefin" : objet avec "categorie" (Clair|Compliqué|Complexe|Pernicieux) et "justification" (1-2 phrases, appuyée sur les données disponibles)
- "ambiguites" : liste de 3-5 formulations ambiguës ou zones d'ombre dans l'énoncé
- "manques" : liste de 3-4 informations clés absentes qui empêchent de bien cerner le problème (ne pas signaler ce qui est déjà fourni par SofIA)
"""
        with st.spinner("Analyse en cours…"):
            raw = claude(prompt)
        data = parse_json(raw)
        if data:
            st.session_state.reformulation = data.get("reformulation", "")
            st.session_state.cynefin = data.get("cynefin", {})
            st.session_state.ambiguites = data.get("ambiguites", [])
            st.session_state.manques = data.get("manques", [])

    # ── Résultats ──
    if st.session_state.reformulation:
        st.divider()
        st.subheader("📊 Analyse de votre problème")

        col_a, col_b = st.columns([3, 1])
        with col_a:
            st.markdown("**🔄 Reformulation proposée**")
            st.info(st.session_state.reformulation)

            if st.session_state.ambiguites:
                st.markdown("**⚠️ Ambiguïtés détectées**")
                for a in st.session_state.ambiguites:
                    st.warning(f"• {a}")

            if st.session_state.manques:
                st.markdown("**❓ Informations manquantes**")
                for m in st.session_state.manques:
                    st.error(f"• {m}")

        with col_b:
            st.markdown("**🧭 Nature du problème**")
            cyn = st.session_state.cynefin
            if isinstance(cyn, dict):
                cat = cyn.get("categorie", "Complexe")
                justif = cyn.get("justification", "")
            else:
                cat, justif = str(cyn), ""
            emoji, desc = CYNEFIN_INFO.get(cat, ("📌", cat))
            st.markdown(f"### {emoji} {cat}")
            st.caption(desc)
            if justif:
                st.caption(f"*{justif}*")

        st.markdown("**✏️ Ajustez la reformulation si nécessaire**")
        edited = st.text_area("Reformulation", value=st.session_state.reformulation, height=100,
                               label_visibility="collapsed")
        st.session_state.reformulation = edited

        st.button("✅ Valider et passer au Miroir des Impossibles →", type="primary", on_click=next_phase)


# ════════════════════════════════════════════════════════════════════════════════
# PHASE 1 — MIROIR DES IMPOSSIBLES
# ════════════════════════════════════════════════════════════════════════════════
elif st.session_state.phase == 1:
    phase_header(
        "2", "Miroir des Impossibles",
        "Avant de cadrer votre sous-problème, un détour nécessaire."
    )

    st.info(f"**Domaine exploré** : {st.session_state.reformulation[:120]}…")

    st.markdown("""
> *"Les experts sont souvent les derniers à croire qu'une rupture est possible dans leur domaine —
> et les premiers à devoir l'intégrer une fois qu'elle arrive."*

Cette phase ne remet pas en cause votre expertise. Elle pose une seule question :
**existe-t-il des technologies ou approches venues d'autres domaines que vous n'avez pas encore considérées ?**
""")

    if st.button("🪞 Générer le Miroir des Impossibles", type="primary",
                 disabled=not API_KEY):
        prompt = f"""L'expert travaille sur ce problème de transition écologique :
"{st.session_state.reformulation}"

Génère 4 cas historiques réels où des experts reconnus d'un domaine ont déclaré qu'une technologie ou approche était impossible — et où ils avaient tort. Choisis des cas variés en termes de domaine (pas uniquement l'énergie), mais tous avec une analogie directe avec le problème décrit.

Pour chaque cas, fournis en JSON valide (aucun markdown) :
- "domaine_origine" : le domaine du cas historique (ex: "Aérospatial", "Médecine", "Informatique"...)
- "ce_que_les_experts_disaient" : la conviction des experts de l'époque en 1 phrase (ton affirmatif, passé)
- "ce_qui_s_est_passe" : ce qui s'est réellement passé en 1-2 phrases (rupture, technologie venue d'ailleurs...)
- "technologie_cle" : la technologie ou approche externe qui a rendu la chose possible
- "analogie_avec_le_probleme" : en 1 phrase, en quoi ce cas éclaire le problème de l'expert (lien direct, pas métaphorique)
- "question_miroir" : une question déstabilisante mais bienveillante que ce cas pose à l'expert sur son propre problème

Réponds UNIQUEMENT avec un JSON valide : {{ "cas": [ ... ] }}"""

        with st.spinner("Recherche de cas analogues…"):
            raw = claude(prompt, max_tokens=2000)
        data = parse_json(raw)
        if data:
            st.session_state.miroir_cas = data.get("cas", [])

    if st.session_state.miroir_cas:
        st.divider()
        st.subheader("📋 Cas historiques — ce que les experts disaient")

        for i, cas in enumerate(st.session_state.miroir_cas):
            with st.expander(
                f"**{cas.get('domaine_origine', '')}** — *\"{cas.get('ce_que_les_experts_disaient', '')}\"*",
                expanded=(i == 0)
            ):
                col_a, col_b = st.columns(2)
                with col_a:
                    st.markdown("**❌ Ce que les experts disaient**")
                    st.error(cas.get("ce_que_les_experts_disaient", ""))
                    st.markdown("**✅ Ce qui s'est passé**")
                    st.success(cas.get("ce_qui_s_est_passe", ""))
                with col_b:
                    st.markdown("**🔑 Technologie clé venue d'ailleurs**")
                    st.info(f"`{cas.get('technologie_cle', '')}`")
                    st.markdown("**🔗 Analogie avec votre problème**")
                    st.warning(cas.get("analogie_avec_le_probleme", ""))

                st.markdown(f"**❓ Question miroir pour vous :** *{cas.get('question_miroir', '')}*")

                # Réaction de l'expert
                reaction = st.text_area(
                    "Votre réaction / ce que ce cas vous inspire :",
                    value=st.session_state.miroir_reactions.get(f"cas_{i}", ""),
                    height=70,
                    key=f"reaction_{i}",
                    placeholder="Ce cas me fait penser que… / Dans mon domaine, l'équivalent serait…"
                )
                st.session_state.miroir_reactions[f"cas_{i}"] = reaction

        st.divider()
        st.subheader("💭 Votre leçon personnelle")
        st.caption(
            "En regardant ces cas ensemble : quelle conviction sur votre problème mérite d'être "
            "mise en question ? Pas pour l'abandonner — pour laisser de la place à l'inattendu."
        )
        lecon = st.text_area(
            "Leçon",
            value=st.session_state.miroir_lecon,
            height=100,
            label_visibility="collapsed",
            placeholder="Ex : Je réalise que je raisonne uniquement à partir des technologies actuelles de captation. Il existe peut-être des approches biologiques ou des matériaux que je n'ai pas explorés…"
        )
        st.session_state.miroir_lecon = lecon

        st.button("✅ Continuer vers le Scout Technologique →", type="primary", on_click=next_phase)


# ════════════════════════════════════════════════════════════════════════════════
# PHASE 2 — SCOUT TECHNOLOGIQUE
# ════════════════════════════════════════════════════════════════════════════════
elif st.session_state.phase == 2:
    phase_header(
        "3", "Scout Technologique",
        "Quelles technologies venues d'autres domaines pourraient changer la donne ?"
    )

    st.info(f"**Problème** : {st.session_state.reformulation[:120]}…")
    if st.session_state.miroir_lecon:
        st.success(f"**Votre leçon du Miroir** : *{st.session_state.miroir_lecon}*")

    st.markdown("""
> Le DARPA finance des projets en imposant une règle : **aucun expert du domaine ne peut être
> chef de projet.** Non par mépris de l'expertise — mais pour forcer les connexions
> avec des technologies que seul un "outsider" oserait proposer.

Cette phase joue ce rôle : identifier des technologies ou méthodes récentes venues d'autres
champs qui pourraient s'appliquer à votre problème.
""")

    if st.button("🔭 Lancer le Scout Technologique", type="primary", disabled=not API_KEY):
        lecon_block = f"\nLeçon que l'expert tire du Miroir des Impossibles : {st.session_state.miroir_lecon}" if st.session_state.miroir_lecon else ""

        prompt = f"""L'expert travaille sur ce problème de transition écologique :
"{st.session_state.reformulation}"{lecon_block}

Joue le rôle d'un "scout technologique" : identifie 5 technologies ou approches récentes (post-2015) venues de domaines *extérieurs* à la transition écologique qui pourraient apporter une solution partielle ou totalement nouvelle à ce problème.

Domaines à explorer prioritairement : IA/ML, biologie synthétique, matériaux avancés, robotique, blockchain/DLT, sciences comportementales, médecine de précision, informatique quantique, impression 3D avancée, satellites/télédétection, edge computing.

Pour chaque technologie, fournis en JSON valide :
- "technologie" : nom précis de la technologie ou approche
- "domaine_origine" : domaine d'où elle vient
- "maturite" : "émergente" (TRL 1-4) | "en développement" (TRL 5-7) | "disponible" (TRL 8-9)
- "description" : 2 phrases — ce que c'est et ce qu'elle permet
- "application_possible" : 2-3 phrases — comment elle pourrait s'appliquer au problème de l'expert. Sois concret et direct.
- "sous_probleme_potentiel" : quel sous-problème spécifique cette technologie permettrait-elle d'aborder ?
- "risque_principal" : 1 phrase — le principal risque ou limite à date
- "exemple_concret" : 1 cas réel d'application (même dans un autre secteur que la transition écologique)

Réponds UNIQUEMENT avec un JSON valide : {{ "technologies": [ ... ] }}"""

        with st.spinner("Exploration en cours…"):
            raw = claude(prompt, max_tokens=2500)
        data = parse_json(raw)
        if data:
            st.session_state.scout_techs = data.get("technologies", [])

    if st.session_state.scout_techs:
        st.divider()
        st.subheader("🔭 Technologies candidates")
        st.caption("Pour chaque technologie : évaluez si elle ouvre un angle nouveau sur votre problème.")

        maturite_colors = {
            "émergente": ("🔴", "Émergente"),
            "en développement": ("🟡", "En développement"),
            "disponible": ("🟢", "Disponible"),
        }

        sel_ids = [t.get("technologie") for t in st.session_state.scout_selections]

        for tech in st.session_state.scout_techs:
            t_name = tech.get("technologie", "")
            mat = tech.get("maturite", "").lower()
            mat_e, mat_l = maturite_colors.get(mat, ("📌", mat))
            is_sel = t_name in sel_ids

            border = "border-left: 4px solid #1a7a4a;" if is_sel else "border-left: 4px solid #ddd;"
            sel_badge = " ✅" if is_sel else ""

            with st.container():
                c1, c2 = st.columns([5, 1])
                with c1:
                    st.markdown(f"<div style='{border} padding-left:12px'>", unsafe_allow_html=True)
                    st.markdown(
                        f"**{t_name}**{sel_badge} — {mat_e} `{mat_l}` — *origine : {tech.get('domaine_origine', '')}*"
                    )
                    st.markdown(tech.get("description", ""))

                    col_app, col_risk = st.columns(2)
                    with col_app:
                        st.markdown("**🎯 Application possible**")
                        st.info(tech.get("application_possible", ""))
                        st.markdown(f"**📦 Sous-problème adressable :** *{tech.get('sous_probleme_potentiel', '')}*")
                    with col_risk:
                        st.markdown("**⚠️ Risque principal**")
                        st.warning(tech.get("risque_principal", ""))
                        st.markdown(f"**💡 Exemple réel :** {tech.get('exemple_concret', '')}")
                    st.markdown("</div>", unsafe_allow_html=True)

                with c2:
                    if is_sel:
                        if st.button("✖ Retirer", key=f"desel_t_{t_name}", use_container_width=True):
                            st.session_state.scout_selections = [
                                t for t in st.session_state.scout_selections
                                if t.get("technologie") != t_name
                            ]
                            st.rerun()
                    else:
                        if st.button("＋ Retenir", key=f"sel_t_{t_name}", use_container_width=True):
                            st.session_state.scout_selections.append(tech)
                            st.rerun()
                st.divider()

        # Synthèse et ouverture
        if st.session_state.scout_selections:
            st.subheader(f"✅ {len(st.session_state.scout_selections)} technologie(s) retenue(s)")
            for t in st.session_state.scout_selections:
                st.markdown(f"- **{t.get('technologie')}** ({t.get('maturite')}) → *{t.get('sous_probleme_potentiel', '')}*")

        st.divider()
        st.subheader("🔓 Comment cela change-t-il votre regard sur le problème ?")
        st.caption(
            "Pas besoin d'avoir choisi une technologie pour répondre. "
            "Même un 'aucune ne s'applique' est une information utile."
        )
        ouvertures = st.text_area(
            "Ouvertures",
            value=st.session_state.scout_ouvertures,
            height=100,
            label_visibility="collapsed",
            placeholder="Ex : Je n'avais pas envisagé que la télédétection satellite pourrait remplacer les capteurs terrain — cela ouvre la possibilité de couvrir 10x plus de sites avec le même budget…"
        )
        st.session_state.scout_ouvertures = ouvertures

        st.button("✅ Continuer vers la Décomposition →", type="primary", on_click=next_phase)


# ════════════════════════════════════════════════════════════════════════════════
# PHASE 3 — DÉCOMPOSITION
# ════════════════════════════════════════════════════════════════════════════════
elif st.session_state.phase == 3:
    phase_header("2", "Décomposition", "Explorer les sous-problèmes possibles et choisir votre cible.")

    st.info(f"**Problème général** : {st.session_state.reformulation}")
    st.markdown("""
> Un problème général peut se décomposer de dizaines de façons selon l'acteur ciblé,
> le levier utilisé et l'horizon temporel. Cette phase vous oblige à **faire un choix**
> plutôt que de tout vouloir résoudre.
""")

    if st.button("🤖 Générer des sous-problèmes candidats", type="primary",
                 disabled=not API_KEY):

        # Contexte SofIA
        sofia_context = ""
        if st.session_state.get("sofia_data"):
            sd = st.session_state.sofia_data
            facts = [ex.get("reponse_sofia", "")[:500] for ex in sd.get("exchanges", [])]
            sofia_context = (
                "\n\nDONNÉES ADEME (issues de SofIA) :\n"
                + "\n---\n".join(facts[:3])
                + "\n\nUtilise ces données pour ancrer les sous-problèmes dans des objectifs réglementaires réels.\n"
            )

        # Contexte Scout Technologique
        scout_context = ""
        if st.session_state.scout_selections:
            techs = st.session_state.scout_selections
            scout_context = (
                "\n\nTECHNOLOGIES IDENTIFIÉES PAR LE SCOUT TECHNOLOGIQUE :\n"
                + "\n".join(
                    f"- {t.get('technologie')} ({t.get('domaine_origine')}) → "
                    f"Sous-problème potentiel : {t.get('sous_probleme_potentiel', '')}"
                    for t in techs
                )
                + "\n\nAu moins 1-2 des 6 sous-problèmes générés DOIVENT intégrer une de ces technologies "
                  "comme levier principal. Signale-le explicitement dans le champ 'tech_scout'.\n"
            )
        elif st.session_state.scout_ouvertures:
            scout_context = (
                f"\n\nOUVERTURES DE L'EXPERT APRÈS LE SCOUT : {st.session_state.scout_ouvertures}\n"
                "Intègre ces ouvertures dans au moins un sous-problème.\n"
            )

        prompt = f"""Problème général (transition écologique) :
"{st.session_state.reformulation}"{sofia_context}{scout_context}

Génère 6 sous-problèmes candidats distincts. Varie les angles :
- Acteur ciblé (entreprise, collectivité, citoyen, filière sectorielle...)
- Levier principal (technique, comportemental, réglementaire, économique, formation, data...)
- Temporalité (court terme <18 mois vs moyen terme 18-36 mois)

Pour chaque sous-problème, fournis :
- "id" : numéro de 1 à 6
- "titre" : titre court accrocheur (max 10 mots)
- "description" : 2 phrases précises. Cite un objectif chiffré ou une réglementation si pertinent.
- "acteur" : acteur principal ciblé
- "levier" : type de levier principal (1 mot)
- "temporalite" : "court" (<18 mois) ou "moyen" (18-36 mois)
- "ambition" : "locale", "sectorielle" ou "systémique"
- "pourquoi_pertinent" : 1 phrase avec données concrètes si disponibles
- "ancrage_ademe" : "oui" si ancré dans des objectifs ADEME, "non" sinon
- "tech_scout" : si ce sous-problème mobilise une technologie identifiée par le Scout, indiquer son nom — sinon null

Réponds UNIQUEMENT en JSON valide avec une clé "sous_problemes" contenant la liste.
"""
        with st.spinner("Génération des sous-problèmes…"):
            raw = claude(prompt, max_tokens=2500)
        data = parse_json(raw)
        if data:
            st.session_state.sub_problems = data.get("sous_problemes", [])

    if st.session_state.sub_problems:
        st.divider()

        sel_list = st.session_state.selected_sub_problems_list
        nb_sel = len(sel_list)
        sel_ids = [sp.get("id") for sp in sel_list]

        # ── Compteur de sélection ──
        if nb_sel == 0:
            st.info("☝️ Sélectionnez **1 ou 2 sous-problèmes** pour continuer. Avec 2 sélections, Claude les synthétisera en un sous-problème affiné.")
        elif nb_sel == 1:
            st.warning(f"**1/2 sélectionné** : *{sel_list[0].get('titre', '')}* — Sélectionnez un 2ᵉ pour une synthèse, ou validez directement.")
        else:
            st.success(f"**2/2 sélectionnés** : *{sel_list[0].get('titre', '')}* + *{sel_list[1].get('titre', '')}*")

        st.subheader("🗂️ Sous-problèmes candidats")

        levier_emojis = {
            "technique": "⚙️", "comportemental": "🧠", "réglementaire": "📋",
            "économique": "💰", "formation": "🎓", "data": "📊", "organisationnel": "🏗️"
        }

        for sp in st.session_state.sub_problems:
            sp_id = sp.get("id")
            is_selected = sp_id in sel_ids
            le = levier_emojis.get(sp.get("levier", "").lower(), "📌")
            te = "⏱️" if sp.get("temporalite") == "court" else "📅"
            ae = {"locale": "📍", "sectorielle": "🏭", "systémique": "🌍"}.get(sp.get("ambition", ""), "")
            ademe_badge = " 🏛️ *ancré ADEME*" if sp.get("ancrage_ademe") == "oui" else ""
            scout_badge = f" 🔭 *{sp.get('tech_scout')}*" if sp.get("tech_scout") else ""

            # Bordure verte si sélectionné
            border_style = "border-left: 4px solid #1a7a4a; padding-left: 12px;" if is_selected else "border-left: 4px solid #e0e0e0; padding-left: 12px;"
            sel_badge = " ✅ **SÉLECTIONNÉ**" if is_selected else ""

            with st.container():
                c1, c2 = st.columns([5, 1])
                with c1:
                    st.markdown(
                        f"<div style='{border_style}'>",
                        unsafe_allow_html=True
                    )
                    st.markdown(
                        f"**{sp_id}. {sp.get('titre', '')}** {le} {te} {ae}{ademe_badge}{scout_badge}{sel_badge}\n\n"
                        f"{sp.get('description', '')}\n\n"
                        f"`Acteur : {sp.get('acteur', '')}` &nbsp;|&nbsp; "
                        f"`Levier : {sp.get('levier', '')}` &nbsp;|&nbsp; "
                        f"`{sp.get('ambition', '')}` &nbsp;|&nbsp; "
                        f"`{sp.get('temporalite', '')} terme`\n\n"
                        f"*{sp.get('pourquoi_pertinent', '')}*"
                    )
                    st.markdown("</div>", unsafe_allow_html=True)

                with c2:
                    if is_selected:
                        if st.button("✖ Retirer", key=f"desel_{sp_id}", use_container_width=True):
                            st.session_state.selected_sub_problems_list = [
                                s for s in sel_list if s.get("id") != sp_id
                            ]
                            st.rerun()
                    else:
                        disabled = nb_sel >= 2
                        btn_label = "＋ Choisir" if nb_sel < 2 else "Max 2"
                        if st.button(btn_label, key=f"sel_{sp_id}",
                                     use_container_width=True, disabled=disabled):
                            st.session_state.selected_sub_problems_list.append(sp)
                            st.rerun()
                st.divider()

    # ── Zone de synthèse / validation ──────────────────────────────────────────
    sel_list = st.session_state.selected_sub_problems_list
    nb_sel = len(sel_list)

    if nb_sel == 2:
        st.subheader("🔀 Synthèse des 2 sous-problèmes")
        st.markdown(
            f"Vous avez sélectionné **{sel_list[0].get('titre', '')}** "
            f"et **{sel_list[1].get('titre', '')}**. "
            "Claude peut les synthétiser en un sous-problème plus précis qui capture "
            "le meilleur des deux angles, ou vous pouvez en choisir un seul."
        )

        col_synth, col_manual = st.columns(2)
        with col_synth:
            if st.button("🤖 Synthétiser les 2 en un sous-problème affiné", type="primary",
                         use_container_width=True, disabled=not API_KEY):
                sp1, sp2 = sel_list[0], sel_list[1]
                prompt_synth = f"""L'expert a sélectionné 2 sous-problèmes complémentaires dans le domaine de la transition écologique. Synthétise-les en un seul sous-problème affiné et cohérent.

PROBLÈME GÉNÉRAL : {st.session_state.reformulation}

SOUS-PROBLÈME A :
- Titre : {sp1.get('titre', '')}
- Description : {sp1.get('description', '')}
- Acteur : {sp1.get('acteur', '')} | Levier : {sp1.get('levier', '')} | {sp1.get('temporalite', '')} terme

SOUS-PROBLÈME B :
- Titre : {sp2.get('titre', '')}
- Description : {sp2.get('description', '')}
- Acteur : {sp2.get('acteur', '')} | Levier : {sp2.get('levier', '')} | {sp2.get('temporalite', '')} terme

Génère en JSON valide (aucun markdown) :
- "titre" : titre du sous-problème synthétisé (max 12 mots)
- "sous_probleme" : formulation complète en 3-4 phrases précises qui intègre les apports des deux angles
- "ce_qui_vient_de_A" : 1 phrase — ce que le sous-problème A apporte à la synthèse
- "ce_qui_vient_de_B" : 1 phrase — ce que le sous-problème B apporte à la synthèse
- "acteur" : acteur(s) principaux ciblés
- "levier" : levier(s) principal/principaux
- "gain_vs_choix_unique" : 1-2 phrases — pourquoi cette synthèse est meilleure que choisir l'un ou l'autre seul
"""
                with st.spinner("Synthèse en cours…"):
                    raw = claude(prompt_synth, max_tokens=1200)
                data = parse_json(raw)
                if data:
                    synth_text = (
                        f"[SYNTHÈSE] {data.get('titre', '')} : {data.get('sous_probleme', '')}"
                    )
                    st.session_state.selected_sub_problem = synth_text
                    st.session_state["_synth_data"] = data
                    st.rerun()

        with col_manual:
            st.caption("Ou choisissez lequel des deux conserver :")
            c1, c2 = st.columns(2)
            with c1:
                if st.button(f"Garder A\n*{sel_list[0].get('titre', '')[:25]}…*",
                             use_container_width=True):
                    sp = sel_list[0]
                    st.session_state.selected_sub_problem = f"{sp.get('titre', '')} : {sp.get('description', '')}"
                    st.session_state.selected_sub_problems_list = [sp]
                    st.rerun()
            with c2:
                if st.button(f"Garder B\n*{sel_list[1].get('titre', '')[:25]}…*",
                             use_container_width=True):
                    sp = sel_list[1]
                    st.session_state.selected_sub_problem = f"{sp.get('titre', '')} : {sp.get('description', '')}"
                    st.session_state.selected_sub_problems_list = [sp]
                    st.rerun()

        # Affichage résultat de synthèse
        if st.session_state.get("_synth_data"):
            sd = st.session_state["_synth_data"]
            st.markdown(f"""
<div class="hmw-box" style="margin-top:16px">
  <p style="color:#a8d8ea; margin:0 0 6px 0; font-size:0.85rem; text-transform:uppercase; letter-spacing:1px">Sous-problème synthétisé</p>
  <h3 style="font-size:1.1rem">{sd.get('titre', '')}</h3>
  <p style="color:#d4f5e9; margin:8px 0 0 0; font-size:0.95rem">{sd.get('sous_probleme', '')}</p>
</div>
""", unsafe_allow_html=True)
            col_a, col_b = st.columns(2)
            with col_a:
                st.caption(f"**Apport de A :** {sd.get('ce_qui_vient_de_A', '')}")
            with col_b:
                st.caption(f"**Apport de B :** {sd.get('ce_qui_vient_de_B', '')}")
            if sd.get("gain_vs_choix_unique"):
                st.info(f"💡 {sd['gain_vs_choix_unique']}")

    elif nb_sel == 1:
        sp = sel_list[0]
        st.session_state.selected_sub_problem = f"{sp.get('titre', '')} : {sp.get('description', '')}"

    # ── Formulation manuelle (toujours disponible) ────────────────────────────
    st.subheader("✍️ Formulation finale — ajustez si nécessaire")
    custom = st.text_area(
        "Votre sous-problème",
        value=st.session_state.selected_sub_problem,
        height=120,
        placeholder="Décrivez précisément le sous-problème que vous souhaitez résoudre...",
        label_visibility="collapsed"
    )
    st.session_state.selected_sub_problem = custom

    if st.session_state.selected_sub_problem.strip():
        st.success(f"**Sous-problème retenu** : {st.session_state.selected_sub_problem}")
        st.button("✅ Valider et décrire mes ressources →", type="primary", on_click=next_phase)


# ════════════════════════════════════════════════════════════════════════════════
# PHASE 4 — RESSOURCES & CONTRAINTES
# ════════════════════════════════════════════════════════════════════════════════
elif st.session_state.phase == 4:
    phase_header("5", "Ressources & Contraintes", "Plus vous êtes précis ici, plus le test de cohérence sera utile.")

    st.info(f"**Sous-problème** : {st.session_state.selected_sub_problem}")

    st.markdown("""
> C'est souvent ici que les projets déraillent : les objectifs sont calibrés sur l'ambition,
> pas sur les moyens réels. Soyez honnête — cet outil ne juge pas, il vous aide.
""")

    r = st.session_state.resources
    c = st.session_state.constraints

    col1, col2 = st.columns(2)

    with col1:
        st.subheader("💰 Budget")
        budget_total      = st.number_input("Budget total (€)", min_value=0, step=10_000,
                                             value=int(r.get("budget", 0)))
        budget_rh         = st.number_input("Dont masse salariale / RH (€)", min_value=0, step=5_000,
                                             value=int(r.get("budget_rh", 0)))
        budget_prestation = st.number_input("Dont prestations externes (€)", min_value=0, step=5_000,
                                             value=int(r.get("budget_prestation", 0)))

        st.subheader("👥 Équipe")
        etp = st.number_input("ETP dédiés au projet", min_value=0.0, step=0.5,
                               value=float(r.get("etp", 0.5)))
        competences = st.multiselect(
            "Compétences disponibles en interne",
            ["Ingénierie / technique", "Data / IA", "Sciences comportementales",
             "Communication / sensibilisation", "Droit / réglementation",
             "Finance / économie", "Sciences naturelles / écologie",
             "Management de projet", "Concertation / participation citoyenne"],
            default=r.get("competences", [])
        )
        competences_manquantes = st.text_area(
            "Compétences manquantes (à recruter ou sous-traiter)",
            value=r.get("competences_manquantes", ""),
            height=80,
            placeholder="Ex: expertise juridique en droit environnemental, data scientist..."
        )

    with col2:
        st.subheader("⏱️ Contraintes temporelles")
        duree_mois = st.slider("Durée totale du projet (mois)", 6, 60,
                               value=int(c.get("duree_mois", 36)))
        jalon_mois = st.slider("Premier jalon / résultats intermédiaires attendus (mois)", 3, 36,
                               value=int(c.get("jalon_mois", 18)))

        st.subheader("🎯 Objectifs attendus")
        objectif_principal = st.text_area(
            "Objectif principal (chiffré si possible)",
            value=c.get("objectif_principal", ""),
            height=80,
            placeholder="Ex: Réduire de 20% les émissions GES du secteur X sur le territoire Y d'ici 2027"
        )
        objectif_secondaire = st.text_area(
            "Objectifs secondaires",
            value=c.get("objectif_secondaire", ""),
            height=60,
            placeholder="Ex: Former 50 entreprises, créer un outil de mesure réplicable..."
        )

        st.subheader("🤝 Partenaires & Historique")
        partenaires_imposes = st.text_area(
            "Partenaires obligatoires (imposés par le financement ou la gouvernance)",
            value=c.get("partenaires_imposes", ""),
            height=60,
            placeholder="Ex: ADEME, Région X, Fédération professionnelle Y..."
        )
        partenaires_potentiels = st.text_area(
            "Partenaires potentiels (à mobiliser)",
            value=c.get("partenaires_potentiels", ""),
            height=60
        )
        deja_tente = st.text_area(
            "Ce qui a déjà été tenté sur ce problème (et pourquoi ça n'a pas suffi)",
            value=c.get("deja_tente", ""),
            height=80,
            placeholder="Ex: Campagne de sensibilisation en 2021 → faible adoption car pas de levier économique associé..."
        )

    st.divider()

    if st.button("✅ Valider et tester la cohérence →", type="primary"):
        st.session_state.resources = {
            "budget": budget_total, "budget_rh": budget_rh,
            "budget_prestation": budget_prestation, "etp": etp,
            "competences": competences, "competences_manquantes": competences_manquantes,
        }
        st.session_state.constraints = {
            "duree_mois": duree_mois, "jalon_mois": jalon_mois,
            "objectif_principal": objectif_principal, "objectif_secondaire": objectif_secondaire,
            "partenaires_imposes": partenaires_imposes, "partenaires_potentiels": partenaires_potentiels,
            "deja_tente": deja_tente,
        }
        next_phase()


# ════════════════════════════════════════════════════════════════════════════════
# PHASE 5 — TEST DE COHÉRENCE
# ════════════════════════════════════════════════════════════════════════════════
elif st.session_state.phase == 5:
    phase_header("6", "Test de Cohérence", "Votre sous-problème est-il réellement à votre portée ?")

    r = st.session_state.resources
    c = st.session_state.constraints

    col1, col2, col3, col4 = st.columns(4)
    col1.metric("💰 Budget total", f"{r.get('budget', 0):,.0f} €")
    col2.metric("👥 ETP", f"{r.get('etp', 0)}")
    col3.metric("📅 Durée", f"{c.get('duree_mois', 0)} mois")
    col4.metric("⏱️ 1er jalon", f"{c.get('jalon_mois', 0)} mois")

    st.info(f"**Sous-problème évalué** : {st.session_state.selected_sub_problem}")

    if st.button("🔬 Lancer le test de cohérence", type="primary",
                 disabled=not API_KEY):
        prompt = f"""Tu évalues la cohérence entre un sous-problème de transition écologique et les ressources disponibles.

SOUS-PROBLÈME : {st.session_state.selected_sub_problem}
PROBLÈME GÉNÉRAL : {st.session_state.reformulation}

RESSOURCES :
- Budget total : {r.get('budget', 0):,}€ (RH : {r.get('budget_rh', 0):,}€ | Prestations : {r.get('budget_prestation', 0):,}€)
- ETP : {r.get('etp', 0)}
- Compétences internes : {', '.join(r.get('competences', [])) or 'Non précisé'}
- Compétences manquantes : {r.get('competences_manquantes', 'Non précisé')}

CONTRAINTES :
- Durée totale : {c.get('duree_mois', 0)} mois
- Premier jalon : {c.get('jalon_mois', 0)} mois
- Objectif principal : {c.get('objectif_principal', 'Non précisé')}
- Partenaires obligatoires : {c.get('partenaires_imposes', 'Aucun')}
- Déjà tenté : {c.get('deja_tente', 'Rien de connu')}

Évalue selon 4 critères, chacun noté de 0 à 25 :

1. "mesurabilite" — Les objectifs sont-ils mesurables avec ces moyens ? Des KPIs clairs sont-ils définissables ?
2. "jalons" — Un résultat concret est-il réaliste dans le délai du premier jalon (apprentissage, prototype, données) ?
3. "competences" — Les compétences disponibles couvrent-elles les leviers nécessaires ? Les manques sont-ils gérables ?
4. "precedents" — Existe-t-il dans la bibliographie ou la pratique des projets comparables à ce niveau de ressources ?

Pour chaque critère :
- "score" : 0 à 25
- "commentaire" : 2-3 phrases d'analyse honnête
- "recommandation" : 1 action corrective concrète si score < 15 (sinon null)

Ajoute aussi :
- "score_global" : somme des 4 scores (0-100)
- "verdict" : "Cohérent" (70-100) | "Ajustements nécessaires" (40-69) | "Recadrage fort recommandé" (0-39)
- "analyse_risques" : 2-3 phrases sur les principaux risques de ce sous-problème avec ces ressources
- "recadrage_suggere" : Si verdict ≠ "Cohérent", propose un sous-problème plus ajusté (sinon null)

Réponds UNIQUEMENT en JSON valide. Aucun markdown, aucun backtick.
"""
        with st.spinner("Analyse de cohérence en cours…"):
            raw = claude(prompt, system=(
                "Tu es un expert rigoureux en évaluation de projets de transition écologique. "
                "Tu donnes des verdicts honnêtes et constructifs. Tu réponds en JSON valide uniquement."
            ), max_tokens=2500)
        data = parse_json(raw)
        if data:
            st.session_state.coherence_score = data.get("score_global", 0)
            st.session_state.coherence_details = data

    # ── Résultats ──
    if st.session_state.coherence_score is not None:
        d = st.session_state.coherence_details
        score = st.session_state.coherence_score
        verdict = d.get("verdict", "")
        emoji = score_emoji(score)

        st.divider()

        # Score global centré
        _, mid, _ = st.columns([1, 2, 1])
        with mid:
            st.markdown(f"""
<div class="score-box">
  <h1 style="font-size:3rem; margin:0">{emoji} {score}<span style="font-size:1.5rem">/100</span></h1>
  <h3 style="margin:8px 0 0 0">Verdict : {verdict}</h3>
</div>
""", unsafe_allow_html=True)

        st.divider()

        # Détail par critère
        st.subheader("📊 Détail par critère")
        criteres = [
            ("mesurabilite",  "📏 Mesurabilité des objectifs"),
            ("jalons",        "📅 Faisabilité du premier jalon"),
            ("competences",   "🧠 Adéquation des compétences"),
            ("precedents",    "📚 Précédents comparables"),
        ]
        for key, label in criteres:
            crit = d.get(key, {})
            if isinstance(crit, dict):
                s = crit.get("score", 0)
                comment = crit.get("commentaire", "")
                reco = crit.get("recommandation", "")
                se = score_emoji(s, max_s=25)
                with st.container():
                    c1, c2 = st.columns([1, 4])
                    with c1:
                        st.metric(label, f"{se} {s}/25")
                    with c2:
                        st.markdown(comment)
                        if reco:
                            st.info(f"💡 {reco}")
                st.divider()

        # Analyse risques
        if d.get("analyse_risques"):
            st.subheader("⚠️ Analyse des risques")
            st.warning(d["analyse_risques"])

        # Recadrage
        recadrage = d.get("recadrage_suggere")
        if recadrage and verdict != "Cohérent":
            st.subheader("🔄 Recadrage suggéré par Claude")
            st.markdown(f"> {recadrage}")
            if st.button("↩️ Adopter ce recadrage (retour Phase 2)"):
                st.session_state.selected_sub_problem = recadrage
                go_phase(1)

        st.button("✅ Valider et passer à la Formulation Finale →", type="primary", on_click=next_phase)


# ════════════════════════════════════════════════════════════════════════════════
# PHASE 6 — FORMULATION FINALE & EXPORT
# ════════════════════════════════════════════════════════════════════════════════
elif st.session_state.phase == 6:
    phase_header("7", "Formulation Finale", "Votre 'bon' sous-problème formulé, documenté et prêt à l'action.")

    score = st.session_state.coherence_score or 0
    emoji = score_emoji(score)

    _, mid, _ = st.columns([1, 2, 1])
    with mid:
        st.metric(f"{emoji} Score de cohérence final", f"{score}/100",
                  delta=st.session_state.coherence_details.get("verdict", ""))

    st.divider()

    if st.button("✨ Générer la formulation finale", type="primary",
                 disabled=not API_KEY):
        r = st.session_state.resources
        c = st.session_state.constraints

        prompt = f"""Synthétise le travail réalisé pour formuler le "bon" sous-problème final.

PROBLÈME GÉNÉRAL : {st.session_state.reformulation}
SOUS-PROBLÈME RETENU : {st.session_state.selected_sub_problem}
SCORE DE COHÉRENCE : {score}/100 — {st.session_state.coherence_details.get("verdict", "")}

RESSOURCES : Budget {r.get('budget', 0):,}€ | {r.get('etp', 0)} ETP | {c.get('duree_mois', 0)} mois
OBJECTIF PRINCIPAL : {c.get('objectif_principal', 'Non précisé')}

Génère en JSON valide (aucun markdown, aucun backtick) avec EXACTEMENT cette structure :

"hmw" : string — Formulation "Comment pourrions-nous..." en 1-2 phrases (qui agit, levier, cible, objectif chiffré, délai)

"formulation_complete" : string — Brief de problème en 3-4 phrases

"indicateurs_succes" : liste de 4-5 objets, chacun avec :
  {{ "kpi": "nom de l'indicateur", "cible": "valeur cible chiffrée", "unite": "unité de mesure", "echeance": "délai" }}

"prochaines_etapes" : liste de 3 objets, chacun avec :
  {{ "action": "intitulé court de l'action", "description": "détail opérationnel en 1-2 phrases", "responsable": "qui porte cette action", "echeance": "ex: J+15 ou Mois 1", "livrable": "livrable concret attendu" }}

"risques_principaux" : liste de 3 objets, chacun avec :
  {{ "risque": "intitulé du risque", "mitigation": "action de mitigation concrète" }}

"questions_restantes" : liste de 2-3 objets, chacun avec :
  {{ "question": "la question à résoudre", "enjeu": "pourquoi c'est important", "moment_resolution": "quand doit-elle être tranchée" }}
"""
        with st.spinner("Formulation en cours…"):
            raw = claude(prompt, max_tokens=2500)
        data = parse_json(raw)
        if data:
            st.session_state.hmw = data.get("hmw", "")
            st.session_state.final_data = data

    if st.session_state.hmw:
        fd = st.session_state.final_data

        # HMW box
        st.markdown(f"""
<div class="hmw-box">
  <p style="color:#a8d8ea; margin:0 0 8px 0; font-size:0.9rem; text-transform:uppercase; letter-spacing:1px">
    Comment pourrions-nous…
  </p>
  <h3>{st.session_state.hmw}</h3>
</div>
""", unsafe_allow_html=True)

        st.markdown(f"**Formulation complète :** {fd.get('formulation_complete', '')}")

        st.divider()

        col1, col2 = st.columns(2)

        with col1:
            # ── KPIs ──
            st.subheader("📏 Indicateurs de succès (KPIs)")
            for kpi in fd.get("indicateurs_succes", []):
                if isinstance(kpi, dict):
                    cible = kpi.get("cible", "")
                    unite = kpi.get("unite", "")
                    echeance = kpi.get("echeance", "")
                    st.markdown(
                        f"**{kpi.get('kpi', '')}**  \n"
                        f"🎯 `{cible} {unite}` — ⏱ {echeance}"
                    )
                else:
                    st.markdown(f"• {kpi}")
                st.divider()

            # ── Actions J+30 ──
            st.subheader("⚡ 3 premières actions (J+30)")
            for i, action in enumerate(fd.get("prochaines_etapes", []), 1):
                if isinstance(action, dict):
                    with st.container():
                        st.markdown(
                            f"**{i}. {action.get('action', '')}** "
                            f"— ⏱ `{action.get('echeance', '')}` "
                            f"— 👤 *{action.get('responsable', '')}*"
                        )
                        if action.get("description"):
                            st.caption(action["description"])
                        if action.get("livrable"):
                            st.info(f"📄 Livrable : {action['livrable']}", icon=None)
                else:
                    st.markdown(f"**{i}.** {action}")
                st.divider()

        with col2:
            # ── Risques ──
            st.subheader("⚠️ Risques & mitigations")
            for risk in fd.get("risques_principaux", []):
                if isinstance(risk, dict):
                    st.warning(f"**{risk.get('risque', '')}**")
                    if risk.get("mitigation"):
                        st.caption(f"💡 {risk['mitigation']}")
                else:
                    st.warning(f"• {risk}")

            # ── Questions restantes ──
            if fd.get("questions_restantes"):
                st.subheader("❓ Questions à trancher en cours de projet")
                for q in fd.get("questions_restantes", []):
                    if isinstance(q, dict):
                        with st.expander(f"❓ {q.get('question', '')}"):
                            if q.get("enjeu"):
                                st.markdown(f"**Enjeu :** {q['enjeu']}")
                            if q.get("moment_resolution"):
                                st.markdown(f"**À trancher :** *{q['moment_resolution']}*")
                    else:
                        st.info(f"• {q}")

        st.divider()

        # ── EXPORT ──────────────────────────────────────────────────
        st.subheader("📥 Exporter votre fiche")

        col_a, col_b = st.columns(2)

        # Export DOCX
        with col_a:
            if st.button("📄 Générer la fiche Word (.docx)", use_container_width=True, type="primary"):
                try:
                    from docx import Document as Doc
                    from docx.shared import Pt, RGBColor
                    from docx.enum.text import WD_ALIGN_PARAGRAPH

                    r = st.session_state.resources
                    c = st.session_state.constraints
                    d = st.session_state.coherence_details

                    document = Doc()

                    # Styles de base
                    style = document.styles["Normal"]
                    style.font.name = "Calibri"
                    style.font.size = Pt(11)

                    # Titre
                    title_p = document.add_heading("Fiche de Problème Stratégique", 0)
                    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    sub_p = document.add_paragraph(
                        f"Générée le {datetime.now().strftime('%d/%m/%Y à %H:%M')}"
                    )
                    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    document.add_paragraph("")

                    # 1. HMW
                    document.add_heading("1. Formulation « Comment Pourrions-Nous »", level=1)
                    p = document.add_paragraph(st.session_state.hmw)
                    p.runs[0].bold = True

                    # 2. Formulation complète
                    document.add_heading("2. Formulation Complète du Sous-Problème", level=1)
                    document.add_paragraph(fd.get("formulation_complete", ""))

                    # 3. Problème général
                    document.add_heading("3. Problème Général d'origine", level=1)
                    document.add_paragraph(st.session_state.reformulation)

                    # 4. Score de cohérence
                    document.add_heading("4. Score de Cohérence", level=1)
                    score_v = st.session_state.coherence_score
                    verdict_v = d.get("verdict", "")
                    document.add_paragraph(f"Score global : {score_v}/100  —  Verdict : {verdict_v}")

                    for key, label in [
                        ("mesurabilite", "Mesurabilité"),
                        ("jalons", "Faisabilité jalons"),
                        ("competences", "Compétences"),
                        ("precedents", "Précédents"),
                    ]:
                        crit = d.get(key, {})
                        if isinstance(crit, dict):
                            document.add_paragraph(
                                f"  • {label} : {crit.get('score', 0)}/25 — {crit.get('commentaire', '')}"
                            )

                    # 5. Ressources
                    document.add_heading("5. Ressources & Contraintes", level=1)
                    table = document.add_table(rows=7, cols=2)
                    table.style = "Table Grid"
                    rows_data = [
                        ("Budget total", f"{r.get('budget', 0):,.0f} €"),
                        ("Masse salariale / RH", f"{r.get('budget_rh', 0):,.0f} €"),
                        ("Prestations externes", f"{r.get('budget_prestation', 0):,.0f} €"),
                        ("ETP dédiés", str(r.get("etp", 0))),
                        ("Durée totale", f"{c.get('duree_mois', 0)} mois"),
                        ("Premier jalon", f"{c.get('jalon_mois', 0)} mois"),
                        ("Partenaires obligatoires", c.get("partenaires_imposes", "Aucun")),
                    ]
                    for i, (k, v) in enumerate(rows_data):
                        table.rows[i].cells[0].text = k
                        table.rows[i].cells[1].text = v

                    # 6. Objectif
                    document.add_heading("6. Objectif Principal", level=1)
                    document.add_paragraph(c.get("objectif_principal", "Non précisé"))
                    if c.get("objectif_secondaire"):
                        document.add_heading("Objectifs secondaires", level=2)
                        document.add_paragraph(c.get("objectif_secondaire"))

                    # 7. KPIs
                    document.add_heading("7. Indicateurs de Succès (KPIs)", level=1)
                    for kpi in fd.get("indicateurs_succes", []):
                        if isinstance(kpi, dict):
                            document.add_paragraph(
                                f"• {kpi.get('kpi', '')} : {kpi.get('cible', '')} {kpi.get('unite', '')} "
                                f"— Échéance : {kpi.get('echeance', '')}"
                            )
                        else:
                            document.add_paragraph(f"• {kpi}")

                    # 8. Actions J+30
                    document.add_heading("8. Prochaines Étapes (J+30)", level=1)
                    for i, action in enumerate(fd.get("prochaines_etapes", []), 1):
                        if isinstance(action, dict):
                            p = document.add_paragraph(style="List Number")
                            p.add_run(f"{action.get('action', '')}").bold = True
                            document.add_paragraph(
                                f"   {action.get('description', '')}\n"
                                f"   Responsable : {action.get('responsable', '')}  |  "
                                f"Échéance : {action.get('echeance', '')}  |  "
                                f"Livrable : {action.get('livrable', '')}"
                            )
                        else:
                            document.add_paragraph(f"{i}. {action}")

                    # 9. Risques
                    document.add_heading("9. Risques Principaux", level=1)
                    for risk in fd.get("risques_principaux", []):
                        if isinstance(risk, dict):
                            document.add_paragraph(
                                f"• {risk.get('risque', '')} → Mitigation : {risk.get('mitigation', '')}"
                            )
                        else:
                            document.add_paragraph(f"• {risk}")

                    # 10. Questions restantes
                    if fd.get("questions_restantes"):
                        document.add_heading("10. Questions à Trancher en Cours de Projet", level=1)
                        for q in fd.get("questions_restantes", []):
                            if isinstance(q, dict):
                                document.add_paragraph(
                                    f"• {q.get('question', '')}\n"
                                    f"   Enjeu : {q.get('enjeu', '')}\n"
                                    f"   À trancher : {q.get('moment_resolution', '')}"
                                )
                            else:
                                document.add_paragraph(f"• {q}")

                    buf = io.BytesIO()
                    document.save(buf)
                    buf.seek(0)

                    st.download_button(
                        label="⬇️ Télécharger la fiche Word",
                        data=buf.getvalue(),
                        file_name=f"fiche_sous_probleme_{datetime.now().strftime('%Y%m%d_%H%M')}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True,
                    )

                except ImportError:
                    st.error("python-docx non installé. Lancez : `pip install python-docx`")
                except Exception as e:
                    st.error(f"Erreur génération DOCX : {e}")

        # Export JSON
        with col_b:
            export = {
                "meta": {"date": datetime.now().isoformat(), "version": "1.0"},
                "probleme_general": st.session_state.reformulation,
                "sous_probleme_retenu": st.session_state.selected_sub_problem,
                "formulation_hmw": st.session_state.hmw,
                "formulation_complete": fd.get("formulation_complete", ""),
                "coherence": {
                    "score": st.session_state.coherence_score,
                    "verdict": st.session_state.coherence_details.get("verdict", ""),
                    "detail": {
                        k: st.session_state.coherence_details.get(k, {})
                        for k in ("mesurabilite", "jalons", "competences", "precedents")
                    },
                },
                "ressources": st.session_state.resources,
                "contraintes": st.session_state.constraints,
                "indicateurs_succes": fd.get("indicateurs_succes", []),
                "prochaines_etapes": fd.get("prochaines_etapes", []),
                "risques": fd.get("risques_principaux", []),
                "questions_restantes": fd.get("questions_restantes", []),
            }
            st.download_button(
                label="⬇️ Exporter en JSON",
                data=json.dumps(export, ensure_ascii=False, indent=2),
                file_name=f"sous_probleme_{datetime.now().strftime('%Y%m%d_%H%M')}.json",
                mime="application/json",
                use_container_width=True,
            )

        st.divider()
        st.success("""
✅ **Félicitations !** Vous avez formulé votre "bon" sous-problème.

Cette fiche est votre référence de projet. Si le score de cohérence est sous 70, revoyez
le périmètre avant de vous lancer. Un bon problème bien formulé, c'est déjà 50% du travail.
        """)
